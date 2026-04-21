import os
import re
import logging
from datetime import datetime
from typing import Optional, Dict, Any, List
import fitz
from docx import Document
from openpyxl import load_workbook

logging.basicConfig(level=logging.INFO, encoding='utf-8')
logger = logging.getLogger(__name__)

DOCS_FOLDER = "documents"
os.makedirs(DOCS_FOLDER, exist_ok=True)

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.xlsx', '.txt', '.png', '.jpg', '.jpeg']
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """Main entry point - process any supported file."""
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.process_pdf(file_path)
        elif ext == '.docx':
            return self.process_docx(file_path)
        elif ext == '.xlsx':
            return self.process_xlsx(file_path)
        elif ext == '.txt':
            return self.process_txt(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self.process_image(file_path)
        else:
            return {"error": f"Unsupported file type: {ext}"}
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF."""
        try:
            doc = fitz.open(file_path)
            text = ""
            pages = len(doc)
            
            for page_num in range(pages):
                page = doc[page_num]
                text += page.get_text() + "\n\n"
            
            metadata = doc.metadata
            doc.close()
            
            extracted_data = self._extract_key_data(text)
            
            return {
                "status": "success",
                "type": "pdf",
                "filename": os.path.basename(file_path),
                "pages": pages,
                "text": text.strip()[:5000],
                "full_text_available": len(text) > 5000,
                "metadata": {
                    "title": metadata.get("title", ""),
                    "author": metadata.get("author", ""),
                    "subject": metadata.get("subject", ""),
                },
                "extracted_data": extracted_data,
                "suggested_name": self._generate_filename(text, extracted_data)
            }
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word document."""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            extracted_data = self._extract_key_data(text)
            
            return {
                "status": "success",
                "type": "docx",
                "filename": os.path.basename(file_path),
                "text": text.strip()[:5000],
                "tables": tables[:5],
                "extracted_data": extracted_data,
                "suggested_name": self._generate_filename(text, extracted_data)
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Extract data from Excel file."""
        try:
            wb = load_workbook(file_path, data_only=True)
            sheets_data = {}
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows.append([str(cell) if cell is not None else "" for cell in row])
                sheets_data[sheet_name] = rows[:100]
            
            wb.close()
            
            all_text = " ".join([
                " ".join(row) for sheet in sheets_data.values() for row in sheet
            ])
            extracted_data = self._extract_key_data(all_text)
            
            return {
                "status": "success",
                "type": "xlsx",
                "filename": os.path.basename(file_path),
                "sheets": list(sheets_data.keys()),
                "preview": sheets_data,
                "extracted_data": extracted_data,
                "suggested_name": self._generate_filename(all_text, extracted_data)
            }
        except Exception as e:
            logger.error(f"XLSX processing error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def process_txt(self, file_path: str) -> Dict[str, Any]:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            extracted_data = self._extract_key_data(text)
            
            return {
                "status": "success",
                "type": "txt",
                "filename": os.path.basename(file_path),
                "text": text.strip()[:5000],
                "extracted_data": extracted_data,
                "suggested_name": self._generate_filename(text, extracted_data)
            }
        except Exception as e:
            logger.error(f"TXT processing error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image - basic extraction (OCR requires pytesseract)."""
        try:
            from PIL import Image
            
            img = Image.open(file_path)
            width, height = img.size
            format_name = img.format
            
            try:
                import pytesseract
                text = pytesseract.image_to_string(img)
                extracted_data = self._extract_key_data(text)
                suggested_name = self._generate_filename(text, extracted_data)
            except ImportError:
                text = "[OCR not available - install pytesseract]"
                extracted_data = {}
                suggested_name = f"image_{datetime.now().strftime('%Y%m%d')}"
            
            return {
                "status": "success",
                "type": "image",
                "filename": os.path.basename(file_path),
                "dimensions": f"{width}x{height}",
                "format": format_name,
                "text": text.strip()[:5000],
                "extracted_data": extracted_data,
                "suggested_name": suggested_name
            }
        except Exception as e:
            logger.error(f"Image processing error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def _extract_key_data(self, text: str) -> Dict[str, Any]:
        """Extract key financial/accounting data from text with intelligent validation."""
        data = {}
        
        # Intelligent email extraction (more precise)
        data["emails"] = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        data["emails"] = [e for e in data["emails"] if not e.endswith('.') and len(e) > 5][:5]
        
        # Intelligent Australian phone extraction
        
        # Australian mobile: 04XX XXX XXX or +61 4XX XXX XXX
        mobile_patterns = [
            r'\+61\s*4[\d\s]{8}',  # +61 4XX XXX XXX
            r'\b04[\d\s]{8}\b',    # 04XX XXX XXX
            r'\b4\d{2}\s?\d{3}\s?\d{3}\b',  # 4X XXX XXX
            r'\(04\)\s*\d{3}\s?\d{3}',  # (04XX) XXX XXX
        ]
        
        # Australian landline: (02) XXXX XXXX or 02 XXXX XXXX
        landline_patterns = [
            r'\+61\s*[2-8]\d\s?\d{4}\s?\d{4}',
            r'\b0[2-8]\d\s?\d{4}\s?\d{4}\b',
            r'\(0[2-8]\)\s*\d{4}\s?\d{4}',
        ]
        
        # Find all potential phones
        all_phones = set()
        for pattern in mobile_patterns + landline_patterns:
            matches = re.findall(pattern, text)
            for m in matches:
                clean = re.sub(r'[^\d+]', '', m)
                if len(clean) >= 10:
                    # Format as +61 X XXX XXX XX
                    if clean.startswith('+61'):
                        formatted = f"+61 {clean[3:]} {clean[6:]} {clean[9:]}"
                    elif clean.startswith('0'):
                        formatted = f"+61 {clean[1:]} {clean[4:]} {clean[7:]}"
                    else:
                        formatted = m
                    all_phones.add(formatted)
        
        data["phones"] = list(all_phones)[:5]
        
        # ABN with validation (11 digits, with 11-digit pattern)
        abn_match = re.search(r'\b\d{2}\s?\d{3}\s?\d{3}\s?\d{3}\b', text)
        if abn_match:
            abn = re.sub(r'\s', '', abn_match.group())
            if len(abn) == 11 and abn.isdigit():
                data["abn"] = abn
                # Validate ABN checksum (basic check)
                data["abn_valid"] = self._validate_abn(abn)
        
        # Amounts - extract with context (total, balance, GST, etc.)
        amount_data = []
        text_lines = text.split('\n')
        for line in text_lines:
            amounts_in_line = re.findall(r'\$[\d,]+\.?\d*', line)
            if amounts_in_line:
                for amt in amounts_in_line:
                    label = line.replace(amt, '').strip()[:30]
                    if any(kw in label.lower() for kw in ['total', 'balance', 'gst', 'subtotal', 'due', 'amount']):
                        amount_data.append({"amount": amt, "label": label})
                    else:
                        amount_data.append({"amount": amt, "label": ""})
        
        data["amounts"] = amount_data[:10]
        
        # Dates - find document dates (invoice date, due date, etc.)
        date_data = []
        for line in text_lines:
            dates_in_line = re.findall(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', line)
            if dates_in_line:
                for dt in dates_in_line:
                    label = line.replace(dt, '').strip()[:30]
                    if any(kw in label.lower() for kw in ['date', 'due', 'issued', 'valid']):
                        date_data.append({"date": dt, "label": label})
                    else:
                        date_data.append({"date": dt, "label": ""})
        
        data["dates"] = date_data[:10]
        
        # Invoice numbers - more flexible pattern
        invoice_patterns = [
            r'(?:invoice|inv|quote|quotation|ref|reference|billing)[#:\s]*([A-Z0-9/-]+)',
            r'(?:invoice|inv)[-#:\s]*(\d+)',
            r'#\s*(\d+)',
        ]
        
        invoice_nums = set()
        for pattern in invoice_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for m in matches:
                if len(m) >= 3 and len(m) <= 20:
                    invoice_nums.add(m.upper())
        
        data["invoice_numbers"] = list(invoice_nums)[:5]
        
        # Document type detection
        text_lower = text.lower()
        doc_type = "document"
        if any(w in text_lower for w in ['tax invoice', 'tax invoice']):
            doc_type = "tax_invoice"
        elif 'invoice' in text_lower:
            doc_type = "invoice"
        elif 'quote' in text_lower or 'quotation' in text_lower:
            doc_type = "quote"
        elif 'receipt' in text_lower:
            doc_type = "receipt"
        elif 'statement' in text_lower:
            doc_type = "statement"
        elif 'bank' in text_lower and 'statement' in text_lower:
            doc_type = "bank_statement"
        elif 'bas' in text_lower:
            doc_type = "bas"
        
        data["document_type"] = doc_type
        
        return data
    
    def _validate_abn(self, abn: str) -> bool:
        """Validate ABN using ATO algorithm."""
        if len(abn) != 11 or not abn.isdigit():
            return False
        # ABN validation weights
        weights = [1, 3, 5, 7, 9, 11, 13, 15, 17, 19]
        abn_digits = [int(d) for d in abn]
        checksum = weights[0] * ((abn_digits[0] - 1) % 10 + 1)
        for i in range(1, 10):
            checksum += weights[i] * abn_digits[i]
        return (checksum % 89) == 0
    
    def _generate_filename(self, text: str, extracted_data: Dict) -> str:
        """Generate a meaningful filename based on extracted content."""
        parts = []
        
        if extracted_data.get("abn"):
            abn = extracted_data["abn"].replace(" ", "")
            parts.append(f"ABN{abn[:8]}")
        
        invoice_nums = extracted_data.get("invoice_numbers", [])
        if invoice_nums:
            parts.append(invoice_nums[0])
        
        date_match = re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', text)
        if date_match:
            date_str = date_match.group().replace("/", "-").replace("\\", "-")
            parts.append(date_str)
        
        keywords = []
        text_lower = text.lower()
        
        if any(w in text_lower for w in ['invoice', 'tax invoice']):
            keywords.append("Invoice")
        elif 'quote' in text_lower or 'quotation' in text_lower:
            keywords.append("Quote")
        elif 'receipt' in text_lower:
            keywords.append("Receipt")
        elif 'statement' in text_lower:
            keywords.append("Statement")
        elif 'bank' in text_lower and 'statement' in text_lower:
            keywords.append("BankStatement")
        
        if keywords:
            parts.append(keywords[0])
        
        if not parts:
            parts.append(f"Doc_{datetime.now().strftime('%Y%m%d')}")
        
        return "_".join(parts)[:50]
    
    def rename_file(self, old_path: str, new_name: str, destination_folder: str = None) -> Dict[str, Any]:
        """Rename/move a file with suggested name to optionally different folder."""
        try:
            base_directory = os.path.dirname(old_path) or DOCS_FOLDER
            ext = os.path.splitext(old_path)[1]
            
            if not new_name.endswith(ext):
                new_name = new_name + ext
            
            # Use destination folder if provided, otherwise use current folder
            target_dir = destination_folder if destination_folder else base_directory
            os.makedirs(target_dir, exist_ok=True)
            
            new_path = os.path.join(target_dir, new_name)
            
            if os.path.exists(new_path):
                base = os.path.splitext(new_name)[0]
                new_path = os.path.join(target_dir, f"{base}_{datetime.now().strftime('%H%M%S')}{ext}")
            
            os.rename(old_path, new_path)
            
            return {
                "status": "success",
                "old_path": old_path,
                "new_path": new_path,
                "new_name": os.path.basename(new_path)
            }
        except Exception as e:
            logger.error(f"Rename error: {e}")
            return {"error": str(e), "status": "failed"}
    
    def list_documents(self, folder: Optional[str] = None) -> List[Dict]:
        """List all documents in the documents folder."""
        target_folder = folder or DOCS_FOLDER
        os.makedirs(target_folder, exist_ok=True)
        
        docs = []
        for filename in os.listdir(target_folder):
            filepath = os.path.join(target_folder, filename)
            if os.path.isfile(filepath):
                ext = os.path.splitext(filename)[1].lower()
                if ext in self.supported_extensions:
                    stat = os.stat(filepath)
                    docs.append({
                        "name": filename,
                        "path": filepath,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "type": ext[1:]
                    })
        
        return sorted(docs, key=lambda x: x["modified"], reverse=True)

processor = DocumentProcessor()

def process_document(file_path: str) -> Dict[str, Any]:
    return processor.process_file(file_path)

def rename_document(old_path: str, new_name: str, destination_folder: str = None) -> Dict[str, Any]:
    return processor.rename_file(old_path, new_name, destination_folder)

def list_documents(folder: Optional[str] = None) -> List[Dict]:
    return processor.list_documents(folder)

def list_folders() -> List[Dict]:
    """List all folders in the documents directory."""
    base = os.path.dirname(DOCS_FOLDER) or "."
    if not os.path.exists(base):
        return [{"name": "documents", "path": "documents"}]
    
    folders = [{"name": "documents", "path": "documents"}]
    for item in os.listdir(base):
        item_path = os.path.join(base, item)
        if os.path.isdir(item_path) and not item.startswith('.'):
            folders.append({"name": item, "path": item})
    
    return folders

def get_instructions_text() -> str:
    """Get active instructions as a string for AI context."""
    try:
        from services import database as db
        instructions = db.get_instructions(active_only=True)
        if not instructions:
            return ""
        return "\n".join([f"- {inst['instruction_text']}" for inst in instructions])
    except Exception as e:
        logger.warning(f"Could not load instructions: {e}")
        return ""


def extract_with_ai(file_path: str) -> Dict[str, Any]:
    """Use Gemini AI for intelligent document extraction."""
    try:
        from services.gemini import get_gemini_client
        client = get_gemini_client()
        if not client:
            return {"error": "AI not available"}
        
        # Get the text first
        result = process_document(file_path)
        text = result.get("text", "")
        
        if not text:
            return {"error": "No text to analyze"}
        
        # Get active instructions
        instructions = get_instructions_text()
        instructions_block = f"\n\nCHERRY'S INSTRUCTIONS (always follow these):\n{instructions}\n" if instructions else ""
        
        prompt = f"""Extract structured data from this accounting document. 

Return ONLY valid JSON (no markdown, no explanation):
{{
  "document_type": "invoice|quote|receipt|statement|bas|contract|other",
  "entity_name": "Business or person name",
  "abn": "ABN if found (11 digits)",
  "phones": ["list of phone numbers with labels"],
  "emails": ["list of emails"],
  "addresses": ["list of addresses"],
  "dates": {{"invoice_date": "", "due_date": "", "issue_date": ""}},
  "amounts": {{"total": "", "gst": "", "subtotal": "", "balance_due": ""}},
  "invoice_number": "Invoice/quote number",
  "line_items": ["list of items/services"],
  "notes": "Any important notes"
}}
{instructions_block}
DOCUMENT TEXT:
{text[:3000]}
"""
        
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt,
            config={'response_mime_type': 'application/json'}
        )
        
        import json
        ai_data = json.loads(response.text.strip())
        
        return {
            "status": "success",
            "ai_extracted": ai_data,
            "regex_extracted": result.get("extracted_data", {}),
            "suggested_name": _generate_ai_filename(ai_data)
        }
    except Exception as e:
        logger.error(f"AI extraction error: {e}")
        return {"error": str(e)}


def _generate_ai_filename(ai_data: Dict) -> str:
    """Generate filename from AI extracted data."""
    parts = []
    
    doc_type = ai_data.get("document_type", "doc")
    entity = ai_data.get("entity_name", "")
    invoice_num = ai_data.get("invoice_number", "")
    dates = ai_data.get("dates", {})
    
    if entity:
        # Get first word of entity name
        parts.append(entity.split()[0][:15])
    
    if invoice_num:
        parts.append(invoice_num[:10])
    
    if dates.get("invoice_date"):
        parts.append(dates["invoice_date"].replace("/", "-"))
    
    if not parts:
        parts.append(doc_type)
    
    return "_".join(parts)[:50]
